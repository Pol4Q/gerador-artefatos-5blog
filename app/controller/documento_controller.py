from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def varrer_e_substituir(elemento, dados):
    for p in elemento.paragraphs:
        for chave, valor in dados.items():
            if chave == "itens":
                continue 
            tag = f"{{{{{chave}}}}}" if not chave.startswith("(") else chave
            
            if tag in p.text:
                p.text = p.text.replace(tag, str(valor))
                # Aplica tamanho 12 em todas as substituições
                for run in p.runs:
                    run.font.size = Pt(12)

    for table in elemento.tables:
        try:
            cab_0 = table.cell(0, 0).text.strip()
            cab_2 = table.cell(0, 2).text.strip()
            
            if "Item" in cab_0 and "Qnt" in cab_2:
                itens = dados.get("itens", [])
                while len(table.rows) > 1:
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)

                for i, item in enumerate(itens):
                    row = table.add_row()
                    row.cells[0].text = str(i + 1)
                    row.cells[1].text = str(item['descricao'])
                    row.cells[2].text = str(item['quantidade'])

                    for idx, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(12)
                            # Centraliza colunas de número e quantidade
                            if idx == 0 or idx == 2:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif idx == 1:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            pass 

        for row in table.rows:
            for cell in row.cells:
                varrer_e_substituir(cell, dados)

def gerar_dfd(dados, nome_arquivo):
    caminho_template = os.path.join("templates", "template_dfd.docx")
    if not os.path.exists(caminho_template):
        raise FileNotFoundError(f"Não encontrei o modelo em: {caminho_template}")

    doc = Document(caminho_template)

    # --- LÓGICA DE MARCAÇÃO (X) ---
    tipo = dados.get("tipo", "")
    mapeamento_tipo = {
        "Material": "() Material.",
        "Serviço não continuado": "() Serviço não continuado;",
        "Serviço continuado SEM dedicação exclusiva de mão de obra": "() Serviço continuado SEM dedicação exclusiva de mão de obra;",
        "Serviço continuado COM dedicação exclusiva de mão de obra": "() Serviço continuado COM dedicação exclusiva de mão de obra;"
    }
    if tipo in mapeamento_tipo:
        dados[mapeamento_tipo[tipo]] = "(X)" + mapeamento_tipo[tipo][2:]
        
    prioridade = dados.get("prioridade", "")
    if prioridade in ["Baixo", "Médio", "Alto"]:
        dados[f"() {prioridade}"] = f"(X) {prioridade}"

    varrer_e_substituir(doc, dados)
    doc.save(nome_arquivo)
    return nome_arquivo